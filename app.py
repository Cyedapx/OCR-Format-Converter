from flask import Flask, render_template, request, send_file
from flask import send_file
from PIL import Image
import io
import os
import pytesseract
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, Frame
from docx import Document


app = Flask(__name__)

UPLOAD_FOLDER = 'uploads/'
ALLOWED_EXTENSIONS = {'jpg', 'jpeg', 'png'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/')
def index():
    return render_template('index.html')


@app.route('/', methods=['POST'])
def upload_file():
    if request.method == 'POST':
        file = request.files['image']
        output_format = request.form['output_format']
        if file and allowed_file(file.filename):
            image = Image.open(file)
            text = pytesseract.image_to_string(image)
            if output_format == 'pdf':
                pdf_buffer = io.BytesIO()
                c = canvas.Canvas(pdf_buffer, pagesize=A4)
                c.setFontSize(12)
                # Create a ParagraphStyle object with the desired properties
                style = getSampleStyleSheet()['Normal']
                style.alignment = 1  # center alignment
                # Create a Frame object with the desired margins and border
                f = Frame(inch, inch, 6*inch, 8*inch, showBoundary=1, leftPadding=0.5*inch, rightPadding=0.5*inch, topPadding=0.5*inch, bottomPadding=0.5*inch)
                # Create a Paragraph object with the text and the style
                p = Paragraph(text, style)
                # Add the Paragraph object to the Frame object
                f.addFromList([p], c)
                # Save the PDF file
                c.save()
                pdf_buffer.seek(0)
                return send_file(pdf_buffer, mimetype='application/pdf', download_name='output.pdf', as_attachment=True)

            elif request.form['output_format'] == 'docx':
                    docx_buffer = io.BytesIO()
                    doc = Document()
                    doc.add_paragraph(text)
                    doc.save(docx_buffer)
                    docx_buffer.seek(0)
                    return send_file(docx_buffer, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', download_name='output.docx', as_attachment=True)
            elif request.form['output_format'] == 'doc':
                    doc_buffer = io.BytesIO()
                    doc = Document()
                    doc.add_paragraph(text)
                    doc.save(doc_buffer)
                    doc_buffer.seek(0)
                    return send_file(doc_buffer, mimetype='application/msword', download_name='output.doc', as_attachment=True)
            else:
                    return "Invalid output_format"
        else:
            return "Invalid output format"
    else:
            return "Invalid file type"
    
    filename = secure_filename(file.filename)
    return render_template('index.html', filename=filename)
if __name__ == '__main__':
    app.run(debug=True)
